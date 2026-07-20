"""
Extractor Agent
===============
Responsible for pulling raw text out of unstructured source documents
(PDF, DOCX, TXT) and turning it into addressable, page/section-tagged
Chunk objects that downstream agents can cite.

Design notes:
- Each chunk gets a stable chunk_id so the Synthesis Agent can cite it and
  the Verifier Agent can look it back up — this is the backbone of the
  "data integrity" requirement in the problem statement.
- Chunking is paragraph-based with a max-length cap, so chunks are small
  enough to cite precisely but large enough to preserve context.
"""
from __future__ import annotations

import re
import uuid
from pathlib import Path

from models import Chunk, ExtractionResult, SourceType

MAX_CHUNK_CHARS = 1200


def _split_into_chunks(text: str, doc_id: str, doc_name: str, page: int | None = None,
                        section: str | None = None) -> list[Chunk]:
    """Split a block of text into paragraph-based chunks under MAX_CHUNK_CHARS."""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[Chunk] = []
    buffer = ""
    for para in paragraphs:
        if len(buffer) + len(para) + 1 <= MAX_CHUNK_CHARS:
            buffer = f"{buffer}\n{para}".strip()
        else:
            if buffer:
                chunks.append(_make_chunk(buffer, doc_id, doc_name, page, section))
            buffer = para
    if buffer:
        chunks.append(_make_chunk(buffer, doc_id, doc_name, page, section))
    return chunks


def _make_chunk(text: str, doc_id: str, doc_name: str, page: int | None, section: str | None) -> Chunk:
    short_id = uuid.uuid4().hex[:8]
    page_part = f"p{page}" if page is not None else "p?"
    return Chunk(
        chunk_id=f"{doc_id}-{page_part}-{short_id}",
        doc_id=doc_id,
        doc_name=doc_name,
        page=page,
        section=section,
        text=text,
    )


def extract_txt(path: Path, doc_id: str) -> ExtractionResult:
    text = path.read_text(encoding="utf-8", errors="ignore")
    chunks = _split_into_chunks(text, doc_id, path.name)
    return ExtractionResult(doc_id=doc_id, doc_name=path.name, source_type=SourceType.TXT, chunks=chunks)


def extract_pdf(path: Path, doc_id: str) -> ExtractionResult:
    warnings: list[str] = []
    chunks: list[Chunk] = []
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise RuntimeError(
            "PyMuPDF is required for PDF extraction. Install with: pip install pymupdf"
        ) from e

    doc = fitz.open(str(path))
    for page_index in range(len(doc)):
        page = doc[page_index]
        text = page.get_text("text")
        if not text.strip():
            warnings.append(f"Page {page_index + 1} had no extractable text (may be scanned/image-only).")
            continue
        chunks.extend(_split_into_chunks(text, doc_id, path.name, page=page_index + 1))
    doc.close()
    return ExtractionResult(doc_id=doc_id, doc_name=path.name, source_type=SourceType.PDF,
                             chunks=chunks, warnings=warnings)


def extract_docx(path: Path, doc_id: str) -> ExtractionResult:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for DOCX extraction. Install with: pip install python-docx"
        ) from e

    document = docx.Document(str(path))
    current_section = "Document"
    buffer_parts: list[str] = []
    chunks: list[Chunk] = []

    def flush():
        nonlocal buffer_parts
        if buffer_parts:
            text = "\n".join(buffer_parts)
            chunks.extend(_split_into_chunks(text, doc_id, path.name, section=current_section))
            buffer_parts = []

    for para in document.paragraphs:
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            flush()
            current_section = para.text.strip() or current_section
            continue
        if para.text.strip():
            buffer_parts.append(para.text.strip())
    flush()

    # Tables: extract as their own chunks so structured data isn't lost
    for t_idx, table in enumerate(document.tables):
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows_text.append(" | ".join(cells))
        table_text = "\n".join(rows_text)
        if table_text.strip():
            chunks.append(_make_chunk(table_text, doc_id, path.name, page=None,
                                       section=f"Table {t_idx + 1}"))

    return ExtractionResult(doc_id=doc_id, doc_name=path.name, source_type=SourceType.DOCX, chunks=chunks)


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_txt,
}


def extract_document(path: str | Path, doc_id: str | None = None) -> ExtractionResult:
    """Entry point: dispatch to the right extractor based on file extension."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"No such file: {path}")
    ext = path.suffix.lower()
    if ext not in EXTRACTORS:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {list(EXTRACTORS)}")
    doc_id = doc_id or path.stem.replace(" ", "_")[:24]
    result = EXTRACTORS[ext](path, doc_id)
    if not result.chunks:
        result.warnings.append("No text could be extracted from this document.")
    return result
